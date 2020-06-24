#!/usr/bin/python3
from sys import argv, exit, stderr
import os

import docx
from docx import Document
from docx.shared import Inches
from docx import exceptions

import csv


#OLD: check_columns = {9:"Datadinascita", 10:"Esperienzacome150oreotutor", 11:"Inipotesiono", 12:"punteggio",13:"graduatoria"}
#NEW:
check_columns = {9:"Datadinascita", 10:"graduatoria", 11:"punteggio"}
dropped_columns = {9,10,11}
field_dest = 3   # NOTE: fields numbered starting from 0
field_student_name = 5   # NOTE: fields numbered starting from 0
field_iscrizione = 6   # NOTE: fields numbered starting from 0
field_voto = 7   # NOTE: fields numbered starting from 0
field_birth = 8   # NOTE: fields numbered starting from 0

def usage(onstream):
    print("\nUsage: %s  file_name.docx" % os.path.basename(argv[0]), file=onstream)
    print("\nExample: ./%s inputs/'Bozza verbaleTUTOR Informatica.docx'" % os.path.basename(argv[0]), file=onstream)
    print("\nCrea i seguenti file:\n  * docenti_active_list.csv\n  * students_LM.csv\n  * students_PHD.csv\n  * TABLE?.csv\n  * le cartelle DEST_? con i file relativi", file=onstream)
    
# THE MAIN PROGRAM:
if len(argv) < 2:
    usage(stderr)
    exit(1)

document = Document(argv[1])
tables = document.tables
print(f"Ci sono {len(document.tables)} tabelle nel documento.")

table_label = {}
i = 0
for paragraph in document.paragraphs:
    if 'CORSO ' in paragraph.text:
        print(paragraph.text)
        i += 1
        table_label[f"TABLE{i}"] = paragraph.text

dest_file = {}        
dest_opened_tables = {}
students = {}
i = 0
for table in tables:
    i += 1
    t_lab = table_label[f"TABLE{i}"]
    print(f"\n\nExtracting data from  TABLE{i} ({t_lab}) ...")

    def render_gen_row(row,tofile):
        col = 0
        if row.cells[field_student_name].paragraphs[0].text != "candidati":
            if row.cells[field_student_name].paragraphs[0].text not in students.keys():
#                students[row.cells[field_student_name].paragraphs[0].text] = [row.cells[field_iscrizione].paragraphs[0].text,row.cells[field_voto].paragraphs[0].text,row.cells[field_birth].paragraphs[0].text,row.cells[field_tutor_exp].paragraphs[0].text,row.cells[field_iter].paragraphs[0].text]
                students[row.cells[field_student_name].paragraphs[0].text] = [row.cells[field_iscrizione].paragraphs[0].text,row.cells[field_voto].paragraphs[0].text,row.cells[field_birth].paragraphs[0].text]
        for cell in row.cells:
            col += 1
            if col not in dropped_columns:
                for paragraph in cell.paragraphs:
                    print(paragraph.text, end=",",file=tofile)
        print(file=tofile)
    
    def render_CDL_and_first_row(i,tofile):
        print(table_label[f"TABLE{i}"],file=tofile)
        col = 0
        for cell in table.rows[0].cells:
            col += 1
            if col in check_columns.keys():
                if not cell.paragraphs[0].text.replace(" ", "").lower()==check_columns[col].replace(" ", "").lower():
                    print(f"ATTENZIONE! La label di col={col} è {cell.paragraphs[0].text} invece che {check_columns[col]} come ti aspettavi",file=tofile)
                    print(f"ATTENZIONE! La label di col={col} è {cell.paragraphs[0].text} invece che {check_columns[col]} come ti aspettavi",file=stderr)
                    assert False
        render_gen_row(table.rows[0],tofile)
        
    with open(f"TABLE{i}.csv","w") as f:
        render_CDL_and_first_row(i,f)
        for row in table.rows[1:]:
            dest=row.cells[field_dest].paragraphs[0].text.replace(" ", "")
            if dest.replace(" ", "") != "": 
                print(dest)
                if not dest in dest_file.keys():
                    if os.path.exists(f"DEST_{dest}"):
                        print(f"ATTENZIONE! La cartella 'DEST_{dest}' esiste già. Devi prima cancellare tutte le cartelle DEST_* dentro la cartella di lavoro.",file=stderr)
                        assert False
                    os.mkdir(f"DEST_{dest}")
                    dest_file[dest]=open(f"DEST_{dest}/table.csv","w")
                    dest_opened_tables[dest] = []
                if i not in dest_opened_tables[dest]:
                    render_CDL_and_first_row(i,dest_file[dest])
                    dest_opened_tables[dest].append(i)

                render_gen_row(row,f)
                render_gen_row(row,dest_file[dest])


with open("students_PHD.csv","w") as f:
    print("candidati,iscrizione,voto ultima laurea,birth",file=f)
    for stud in students.keys():
#        print(f"{stud}{students[stud]}")
        if students[stud][0][0:4] == "DOTT":
            print(stud,end=",",file=f)
            for field in students[stud]:
                print(field,end=",",file=f)
            print(file=f)
    
with open("students_LM.csv","w") as f:
    print("candidati,iscrizione,voto ultima laurea,birth",file=f)
    for stud in students.keys():
#        print(f"{stud}{students[stud]}")
        if students[stud][0][0:4] != "DOTT":
            print(stud,end=",",file=f)
            for field in students[stud]:
                print(field,end=",",file=f)
            print(file=f)

    
with open("docenti_active_list.csv","w") as f:
    for dest in dest_file.keys():
        print(dest,file=f)

