#!/usr/bin/python3
from sys import argv, exit, stderr
import os

import docx
from docx import Document
from docx.shared import Inches
from docx import exceptions
from docx.shared import Pt

import csv

def usage():
    print("\nUsage: %s  file_name.csv  [title_of_table]" % os.path.basename(argv[0]), file=stderr, )
    exit(1)

# THE MAIN PROGRAM:
if len(argv) < 2:
    usage()
    exit(1)

doc = Document()
if len(argv) >= 3:
    doc.add_heading(argv[2])
    #doc.add_paragraph(argv[2])

with open(argv[1], newline='') as f:
    csv_reader = csv.reader(f) 

    csv_headers = next(csv_reader)
    csv_cols = len(csv_headers)

    table = doc.add_table(rows=2, cols=csv_cols)
    
    hdr_cells = table.rows[0].cells

    for i in range(csv_cols):
        hdr_cells[i].text = csv_headers[i]

    for row in csv_reader:
        row_cells = table.add_row().cells
        for i in range(csv_cols):
            row_cells[i].text = row[i]

doc.add_page_break()
doc.save(f"{argv[1].split('.')[0]}.docx")
